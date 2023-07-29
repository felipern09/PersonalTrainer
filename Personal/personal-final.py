from tkinter import *
import tkinter.filedialog
from tkinter import ttk, messagebox, Text
import pandas as pd
from datetime import datetime
from sqlalchemy.orm import sessionmaker
from models import Personal, Aulas, Valores, engine, Usuario
import os
import docx
import win32com.client as client
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from num2words import num2words
import time as t
from openpyxl import Workbook
from openpyxl import load_workbook as l_w
from email.mime.base import MIMEBase
from email import encoders
from dateutil.relativedelta import relativedelta

root = Tk()
root.title("Calcular valores personal")
img = PhotoImage(file='./static/imgs/Icone.png')
root.iconphoto(False, img)
root.geometry('560x400')
root.columnconfigure(0, weight=5)
root.rowconfigure(0, weight=5)

for child in root.winfo_children():
    child.grid_configure(padx=1, pady=3)

my_notebook = ttk.Notebook(root)
my_notebook.pack()

# Variables
nomesaulas = ['Simples 1h', 'Simples 2h', 'Dupla 1h', 'Dupla 2h', 'Tripla 1h', 'Tripla 2h']
hoje = datetime.today()
Session = sessionmaker(bind=engine)
session = Session()
if session.query(Aulas).filter_by(mes=f'{hoje.month}/{hoje.year}').first():
    ultima_plan = f'{hoje.month}/{hoje.year}'
else:
    ultima_plan = f'{hoje.month - 1}/{hoje.year}'

referencia = ['1/2023', '2/2023', '3/2023', '4/2023', '5/2023', '6/2023', '7/2023', '8/2023', '9/2023', '10/2023',
              '11/2023', '12/2023', '1/2024', '2/2024', '3/2024', '4/2024', '5/2024', '6/2024', '7/2024', '8/2024',
              '9/2024', '10/2024', '11/2024', '12/2024', '1/2025', '2/2025', '3/2025', '4/2025', '5/2025', '6/2025',
              '7/2025', '8/2025', '9/2025', '10/2025', '11/2025', '12/2025', '1/2026', '2/2026', '3/2026', '4/2026',
              '5/2026', '6/2026', '7/2026', '8/2026', '9/2026', '10/2026', '11/2026', '12/2026', '1/2027', '2/2027',
              '3/2027', '4/2027', '5/2027', '6/2027', '7/2027', '8/2027', '9/2027', '10/2027', '11/2027', '12/2027',
              '1/2028', '2/2028', '3/2028', '4/2028', '5/2028', '6/2028', '7/2028', '8/2028', '9/2028', '10/2028',
              '11/2028', '12/2028', '1/2029', '2/2029', '3/2029', '4/2029', '5/2029', '6/2029', '7/2029', '8/2029',
              '9/2029', '10/2029', '11/2029', '12/2029', '1/2030', '2/2030', '3/2030', '4/2030', '5/2030', '6/2030',
              '7/2030', '8/2030', '9/2030', '10/2030', '11/2030', '12/2030', '1/2031', '2/2031', '3/2031', '4/2031',
              '5/2031', '6/2031', '7/2031', '8/2031', '9/2031', '10/2031', '11/2031', '12/2031', '1/2032', '2/2032',
              '3/2032', '4/2032', '5/2032', '6/2032', '7/2032', '8/2032', '9/2032', '10/2032', '11/2032', '12/2032',
              '1/2033', '2/2033', '3/2033', '4/2033', '5/2033', '6/2033', '7/2033', '8/2033', '9/2033', '10/2033',
              '11/2033', '12/2033']


def cadastrarpersonal(nome, email, tel, tipo):
    if tipo == 1:
        tipo = 'Interno'
    else:
        tipo = 'Externo'
    pessoa = Personal(nome=str(nome), email=str(email), whatsapp=str(tel), tipo_personal=str(tipo), status='Ativo')
    session.add(pessoa)
    session.commit()


def pesquisaintext(tipo):
    personal = session.query(Personal).filter_by(tipo_personal=tipo).all()
    return int(len(personal))


def pesquisaatinat(status):
    personal = session.query(Personal).filter_by(status=status).all()
    pess=[]
    for item in personal:
        pess.append(item.nome)
        pess.sort()
    tupla = tuple(pess)
    return tupla


def pesquisatodos():
    personal = session.query(Personal).all()
    pess=[]
    for item in personal:
        pess.append(item.nome)
        pess.sort()
    tupla = tuple(pess)
    return tupla


todos = pesquisatodos()
values = pesquisaatinat('Ativo')
intern = int(pesquisaintext('Interno'))
extern = int(pesquisaintext('Externo'))
pers = intern+extern


def cadastrarusuario(nome, email, senha, servidor, porta, caminhoassinatura):
    existente = session.query(Usuario).filter_by(id=1).first()
    if existente:
        existente.nome = nome
        existente.email = email
        existente.senha = senha
        existente.servidor = servidor
        existente.porta = porta
        existente.assinatura = caminhoassinatura
        session.commit()
    else:
        pessoa = Usuario(nome=nome, email=email, senha=senha, servidor=servidor,
                         porta=porta, assinatura=caminhoassinatura)
        session.add(pessoa)
        session.commit()


def totalaulas():
    aulas = session.query(Aulas).filter_by(mes=f'{hoje.month}/{hoje.year}').all()
    soma = 0
    for aula in aulas:
        soma += aula.simples1 + aula.simples2 + aula.dupla1 + aula.dupla2 + aula.tripla1 + aula.tripla2
    return soma


def totalhoras():
    aulas = session.query(Aulas).filter_by(mes=f'{hoje.month}/{hoje.year}').all()
    soma = 0
    for aula in aulas:
        soma += aula.simples1 + (aula.simples2*2) + aula.dupla1 + (aula.dupla2*2) + aula.tripla1 + (aula.tripla2*2)
    return soma


def totalreceita():
    aulas = session.query(Aulas).filter_by(mes=f'{hoje.month}/{hoje.year}').all()
    s = 0
    for aula in aulas:
        s += aula.valortotalemdia
    soma = 'R$ {:,.2f}'.format(s).replace(',', '_').replace('.', ',').replace('_', '.')
    return soma


def verificar(nome):
    pers = session.query(Personal).filter_by(nome=nome).first()
    aulas = session.query(Aulas).filter_by(personal=pers.id).filter_by(mes=f'{hoje.month}/{hoje.year}').first()
    simples1 = aulas.simples1
    simples2 = aulas.simples2
    dupla1 = aulas.dupla1
    dupla2 = aulas.dupla2
    tripla1 = aulas.tripla1
    tripla2 = aulas.tripla2
    return simples1, simples2, dupla1, dupla2, tripla1, tripla2


def quantidade():
    qtidade = len(session.query(Aulas).filter_by(mes=f"{hoje.month}/{hoje.year}").all())
    return qtidade


def mandar_email():
    personais = session.query(Personal).filter_by(status='Ativo').all()
    usuario = session.query(Usuario).filter_by(id=1).first()
    if usuario:
        carregar = Toplevel(root)
        carregar.geometry('300x100')
        carregar.title('Enviando e-mails...')
        imgs = PhotoImage(file='./static/imgs/Icone.png')
        carregar.iconphoto(False, imgs)
        ttk.Label(carregar, text='Os e-mails estão sendo enviados...').grid(column=1, row=1, padx=20)
        email_remetente = usuario.email
        senha = usuario.senha
        # set up the SMTP server
        s = smtplib.SMTP(host=usuario.servidor, port=usuario.porta)
        s.starttls()
        s.login(email_remetente, senha)

        for personal in personais:
            aulas = session.query(Aulas).filter_by(personal=personal.id)\
                .filter_by(mes=str(hoje.month)+'/'+str(hoje.year)).first()
            if aulas:
                simples1 = aulas.simples1
                simples2 = aulas.simples2
                dupla1 = aulas.dupla1
                dupla2 = aulas.dupla2
                tripla1 = aulas.tripla1
                tripla2 = aulas.tripla2
                if hoje.day <= 10:
                    valor_total = round(aulas.valortotalemdia, 2)
                else:
                    valor_total = round(aulas.valortotalatraso, 2)
                msg = MIMEMultipart()
                message = f'''
                Olá, {str(personal.nome).title().split(sep=' ')[0]}!\n
                \n
                Segue seu relatório de aulas de personal a serem pagas para Cia Athletica:\n
                \n
                {simples1} aulas simples de 1h.\n
                {simples2} aulas simples de 2h.\n
                {dupla1} aulas duplas de 1h.\n
                {dupla2} aulas duplas de 2h.\n
                {tripla1} aula tripla de 1h.\n
                {tripla2} aula tripla de 2h.\n
                Total de {simples1+simples2+dupla1+dupla2+tripla1+tripla2} aulas.\n
                Total de taxa: R$ {str("%.2f" % round(valor_total, 2)).replace('.',',')}.\n
                \n
                Seguem dados para o pagamento:\n
                \n
                PIX: 03732305000186
                Banco: Itaú
                Agência: 6205
                C/C: 01588-3\n
                Assim que o pagamento for feito, favor responder esse e-mail com o comprovante bancário.
                \n
                Atenciosamente,\n
                Marcelo Gonçalves
                '''
                # setup the parameters of the message
                msg['From'] = email_remetente
                msg['To'] = str(personal.email).split(sep=' ')[0]
                msg['Subject'] = "Aulas de personal - Cia Athletica"
                msg.attach(MIMEText(message, 'plain', _charset='utf-8'))

                # Anexo PNG
                arquivo_png = usuario.assinatura
                with open(arquivo_png, 'rb') as img_file:
                    imagem = MIMEImage(img_file.read())
                    msg.attach(imagem)

                s.send_message(msg)
                del msg
        s.quit()
        carregar.destroy()
    else:
        tkinter.messagebox.showinfo(title='Cadastrar e-mail remetente!', message='Antes de enviar os e-mails, cadastre '
                                                                                 'o e-mail do remetente!')


def gerar_cobranca(caminho):
    os.rename(caminho, f'.\\Relatorios\\{hoje.year}\\{str(hoje.month).zfill(2)}\\Personal.xls')
    excel = client.Dispatch('excel.application')
    file = os.path.abspath(f'.\\Relatorios\\{hoje.year}\\{str(hoje.month).zfill(2)}\\Personal.xls')
    filename, fileextension = os.path.splitext(file)
    wb = excel.Workbooks.Open(file)
    wb.SaveAs(filename, 51)
    wb.Close()
    valor = session.query(Valores).filter_by(id=1).first()
    plan = pd.read_excel(filename+'.xlsx')
    personal = []
    for item in plan['Personal Trainer']:
        personal.append(item)
        pers = list(set(personal))
        pers.sort()
    plan = plan.rename(columns={'Personal Trainer': 'Personal'})
    for pessoa in pers:
        plan2 = plan[plan.Personal == pessoa]
        if plan2[(plan2['Serviço'] == 'AULA DE PERSONAL') & (plan2['Duração (horas)'] == 1) & (plan2['Cortesia'] == 'N')].empty:
            simples1 = 0
        else:
            simples1 = int(plan2[(plan2['Serviço'] == 'AULA DE PERSONAL') & (plan2['Duração (horas)'] == 1) & (plan2['Cortesia'] == 'N')].iloc[0, 2])

        if plan2[(plan2['Serviço'] == 'AULA DE PERSONAL') & (plan2['Duração (horas)'] == 2) & (plan2['Cortesia'] == 'N')].empty:
            simples2 = 0
        else:
            simples2 = int(plan2[(plan2['Serviço'] == 'AULA DE PERSONAL') & (plan2['Duração (horas)'] == 2) & (plan2['Cortesia'] == 'N')].iloc[0, 2])

        if plan2[(plan2['Serviço'] == 'AULA DE PERSONAL 2 PESSOAS') & (plan2['Duração (horas)'] == 1) & (plan2['Cortesia'] == 'N')].empty:
            dupla1 = 0
        else:
            dupla1 = int(plan2[(plan2['Serviço'] == 'AULA DE PERSONAL 2 PESSOAS') & (plan2['Duração (horas)'] == 1) & (plan2['Cortesia'] == 'N')].iloc[0, 2])

        if plan2[(plan2['Serviço'] == 'AULA DE PERSONAL 2 PESSOAS') & (plan2['Duração (horas)'] == 2) & (plan2['Cortesia'] == 'N')].empty:
            dupla2 = 0
        else:
            dupla2 = int(plan2[(plan2['Serviço'] == 'AULA DE PERSONAL 2 PESSOAS') & (plan2['Duração (horas)'] == 2) & (plan2['Cortesia'] == 'N')].iloc[0, 2])

        if plan2[(plan2['Serviço'] == 'AULA DE PERSONAL 3 PESSOAS') & (plan2['Duração (horas)'] == 1) & (plan2['Cortesia'] == 'N')].empty:
            tripla1 = 0
        else:
            tripla1 = int(plan2[(plan2['Serviço'] == 'AULA DE PERSONAL 3 PESSOAS') & (plan2['Duração (horas)'] == 1) & (plan2['Cortesia'] == 'N')].iloc[0, 2])

        if plan2[(plan2['Serviço'] == 'AULA DE PERSONAL 3 PESSOAS') & (plan2['Duração (horas)'] == 2) & (plan2['Cortesia'] == 'N')].empty:
            tripla2 = 0
        else:
            tripla2 = int(plan2[(plan2['Serviço'] == 'AULA DE PERSONAL 3 PESSOAS') & (plan2['Duração (horas)'] == 2) & (plan2['Cortesia'] == 'N')].iloc[0, 2])

        if simples1 + simples2 <= 10:
            internosimples = valor.internosimples1a10
        else:
            if simples1 + simples2 <= 30:
                internosimples = valor.internosimples11a30
            else:
                if simples1 + simples2 <= 50:
                    internosimples = valor.internosimples31a50
                else:
                    if simples1 + simples2 <= 100:
                        internosimples = valor.internosimples51a100
                    else:
                        if simples1 + simples2 <= 120:
                            internosimples = valor.internosimples101a120
                        else:
                            internosimples = valor.internosimplesacima120

        if dupla1 + dupla2 <= 60:
            internodupla = valor.internodupla1a60
        else:
            if dupla1 + dupla2 <= 119:
                internodupla = valor.internodupla61a119
            else:
                internodupla = valor.internoduplaacima119

        if tripla1 + tripla2 <= 60:
            internotripla = valor.internotripla1a60
        else:
            if tripla1 + tripla2 <= 119:
                internotripla = valor.internotripla61a119
            else:
                internotripla = valor.internotriplaacima119

        if simples1 + simples2 <= 10:
            externosimples = valor.externosimples1a10
        else:
            if simples1 + simples2 <= 30:
                externosimples = valor.externosimples11a30
            else:
                if simples1 + simples2 <= 50:
                    externosimples = valor.externosimples31a50
                else:
                    if simples1 + simples2 <= 100:
                        externosimples = valor.externosimples51a100
                    else:
                        if simples1 + simples2 <= 120:
                            externosimples = valor.externosimples101a120
                        else:
                            externosimples = valor.externosimplesacima120

        if dupla1 + dupla2 <= 60:
            externodupla = valor.externodupla1a60
        else:
            if dupla1 + dupla2 <= 119:
                externodupla = valor.externodupla61a119
            else:
                externodupla = valor.externoduplaacima119

        if tripla1 + tripla2 <= 60:
            externotripla = valor.externotripla1a60
        else:
            if tripla1 + tripla2 <= 119:
                externotripla = valor.externotripla61a119
            else:
                externotripla = valor.externotriplaacima119

        internosimplesnodesc = valor.internosimplesnodesc
        internoduplanodesc = valor.internoduplanodesc
        internotriplanodesc = valor.internotriplanodesc
        externosimplesnodesc = valor.externosimplesnodesc
        externoduplanodesc = valor.externoduplanodesc
        externotriplanodesc = valor.externotriplanodesc

        if session.query(Personal).filter_by(nome=str(pessoa)).first() is None:
            result = messagebox.askquestion(f'Personal {str(pessoa).title().split(sep=" " )[0]} é interno?', f'Personal {str(pessoa).title()} está sendo adicionado no banco de dados do sistema.\nEsse personal é INTERNO?')
            if result == 'yes':
                adicionar = Personal(nome=str(pessoa), email='marcelo.goncalves@ciaathletica.com.br', tipo_personal='Interno', status='Ativo')
                session.add(adicionar)
                session.commit()
            else:
                adicionar = Personal(nome=str(pessoa), email='marcelo.goncalves@ciaathletica.com.br', tipo_personal='Externo', status='Ativo')
                session.add(adicionar)
                session.commit()
            prof = session.query(Personal).filter_by(nome=str(pessoa)).first()
            prof.status = 'Ativo'
            if prof.tipo_personal == 'Interno':
                valor_simples = round((simples1 * internosimples) + (simples2 * 2 * internosimples), 2)
                valor_dupla = round((dupla1 * internodupla) + (dupla2 * 2 * internodupla), 2)
                valor_tripla = round((tripla1 * internotripla) + (tripla2 * 2 * internotripla), 2)
                valor_total = round(valor_simples + valor_dupla + valor_tripla, 2)
            else:
                valor_simples = round((simples1 * externosimples) + (simples2 * 2 * externosimples), 2)
                valor_dupla = round((dupla1 * externodupla) + (dupla2 * 2 * externodupla), 2)
                valor_tripla = round((tripla1 * externotripla) + (tripla2 * 2 * externotripla), 2)
                valor_total = round(valor_simples + valor_dupla + valor_tripla, 2)

            if prof.tipo_personal == 'Interno':
                valor_simplesnd = round((simples1 * internosimplesnodesc) + (simples2 * 2 * internosimplesnodesc), 2)
                valor_dupland = round((dupla1 * internoduplanodesc) + (dupla2 * 2 * internoduplanodesc), 2)
                valor_tripland = round((tripla1 * internotriplanodesc) + (tripla2 * 2 * internotriplanodesc), 2)
                valor_totalnd = round(valor_simplesnd + valor_dupland + valor_tripland, 2)
            else:
                valor_simplesnd = round((simples1 * externosimplesnodesc) + (simples2 * 2 * externosimplesnodesc), 2)
                valor_dupland = round((dupla1 * externoduplanodesc) + (dupla2 * 2 * externoduplanodesc), 2)
                valor_tripland = round((tripla1 * externotriplanodesc) + (tripla2 * 2 * externotriplanodesc), 2)
                valor_totalnd = round(valor_simplesnd + valor_dupland + valor_tripland, 2)
            mespassado = hoje - relativedelta(months=1)
            mesp = session.query(Aulas).filter_by(mes=f'{mespassado.month}/{mespassado.year}').filter_by(personal=prof.id).first()
            if mesp:
                creditomesp = mesp.credito
                debitomesp = mesp.debito
            else:
                creditomesp = 0
                debitomesp = 0
            aulas = Aulas(
                personal=prof.id, mes=str(hoje.month)+'/'+str(hoje.year), simples1=simples1,
                simples2=simples2, dupla1=dupla1, dupla2=dupla2, tripla1=tripla1, tripla2=tripla2,
                valortotalemdia=valor_total - creditomesp + debitomesp,
                valortotalatraso=valor_totalnd - creditomesp + debitomesp
            )
            session.add(aulas)
            session.commit()
        else:
            prof = session.query(Personal).filter_by(nome=str(pessoa)).first()
            prof.status = 'Ativo'
            if prof.tipo_personal == 'Interno':
                valor_simples = round((simples1 * internosimples) + (simples2 * 2 * internosimples), 2)
                valor_dupla = round((dupla1 * internodupla) + (dupla2 * 2 * internodupla), 2)
                valor_tripla = round((tripla1 * internotripla) + (tripla2 * 2 * internotripla), 2)
                valor_total = round(valor_simples + valor_dupla + valor_tripla, 2)
            else:
                valor_simples = round((simples1 * externosimples) + (simples2 * 2 * externosimples), 2)
                valor_dupla = round((dupla1 * externodupla) + (dupla2 * 2 * externodupla), 2)
                valor_tripla = round((tripla1 * externotripla) + (tripla2 * 2 * externotripla), 2)
                valor_total = round(valor_simples + valor_dupla + valor_tripla, 2)

            if prof.tipo_personal == 'Interno':
                valor_simplesnd = round((simples1 * internosimplesnodesc) + (simples2 * 2 * internosimplesnodesc), 2)
                valor_dupland = round((dupla1 * internoduplanodesc) + (dupla2 * 2 * internoduplanodesc), 2)
                valor_tripland = round((tripla1 * internotriplanodesc) + (tripla2 * 2 * internotriplanodesc), 2)
                valor_totalnd = round(valor_simplesnd + valor_dupland + valor_tripland, 2)
            else:
                valor_simplesnd = round((simples1 * externosimplesnodesc) + (simples2 * 2 * externosimplesnodesc), 2)
                valor_dupland = round((dupla1 * externoduplanodesc) + (dupla2 * 2 * externoduplanodesc), 2)
                valor_tripland = round((tripla1 * externotriplanodesc) + (tripla2 * 2 * externotriplanodesc), 2)
                valor_totalnd = round(valor_simplesnd + valor_dupland + valor_tripland, 2)
            mespassado = hoje - relativedelta(months=1)
            mesp = session.query(Aulas).filter_by(mes=f'{mespassado.month}/{mespassado.year}').filter_by(personal=prof.id).first()
            if mesp:
                creditomesp = mesp.credito
                debitomesp = mesp.debito
            else:
                creditomesp = 0
                debitomesp = 0

            jaexiste = session.query(Aulas).filter_by(mes=f'{hoje.month}/{hoje.year}').filter_by(personal=prof.id).first()
            if jaexiste:
                jaexiste.personal = prof.id
                jaexiste.mes = f'{hoje.month}/{hoje.year}'
                jaexiste.simples1 = simples1
                jaexiste.simples2 = simples2
                jaexiste.dupla1 = dupla1
                jaexiste.dupla2 = dupla2
                jaexiste.tripla1 = tripla1
                jaexiste.tripla2 = tripla2
                jaexiste.valortotalemdia = valor_total - creditomesp + debitomesp
                jaexiste.valortotalatraso = valor_totalnd - creditomesp + debitomesp
                session.commit()
            else:
                aulas = Aulas(personal=prof.id, mes=str(hoje.month)+'/'+str(hoje.year), simples1=simples1,
                      simples2=simples2, dupla1=dupla1, dupla2=dupla2, tripla1=tripla1, tripla2=tripla2,
                      valortotalemdia=valor_total - creditomesp + debitomesp, valortotalatraso=valor_totalnd - creditomesp + debitomesp)
                session.add(aulas)
                session.commit()
            session.commit()


def ajuste(nome, nomeaulad='', numeroaulasd=0, descvalor=0, nomeaulaa='', numeroaulasa=0, acresvalor=0):
    valor = session.query(Valores).filter_by(id=1).first()
    if numeroaulasd == '':
        numeroaulasd = 0
    if descvalor == '':
        descvalor = 0
    if numeroaulasa == '':
        numeroaulasa = 0
    if acresvalor == '':
        acresvalor = 0
    if nome:
        prof = session.query(Personal).filter_by(nome=nome).first()
        aula = session.query(Aulas).filter_by(personal=prof.id).filter_by(mes=f'{hoje.month}/{hoje.year}').first()
        if nomeaulad or nomeaulaa != '':
            if nomeaulad == 'Simples 1h':
                aula.simples1 = aula.simples1 - float(str(numeroaulasd).replace(',', '.'))
            if nomeaulad == 'Simples 2h':
                aula.simples2 = aula.simples2 - float(str(numeroaulasd).replace(',', '.'))
            if nomeaulad == 'Dupla 1h':
                aula.dupla1 = aula.dupla1 - float(str(numeroaulasd).replace(',', '.'))
            if nomeaulad == 'Dupla 2h':
                aula.dupla2 = aula.dupla2 - float(str(numeroaulasd).replace(',', '.'))
            if nomeaulad == 'Tripla 1h':
                aula.tripla1 = aula.tripla1 - float(str(numeroaulasd).replace(',', '.'))
            if nomeaulad == 'Tripla 2h':
                aula.tripla2 = aula.tripla2 - float(str(numeroaulasd).replace(',', '.'))
            if nomeaulaa == 'Simples 1h':
                aula.simples1 = aula.simples1 + float(str(numeroaulasa).replace(',', '.'))
            if nomeaulaa == 'Simples 2h':
                aula.simples2 = aula.simples2 + float(str(numeroaulasa).replace(',', '.'))
            if nomeaulaa == 'Dupla 1h':
                aula.dupla1 = aula.dupla1 + float(str(numeroaulasa).replace(',', '.'))
            if nomeaulaa == 'Dupla 2h':
                aula.dupla2 = aula.dupla2 + float(str(numeroaulasa).replace(',', '.'))
            if nomeaulaa == 'Tripla 1h':
                aula.tripla1 = aula.tripla1 + float(str(numeroaulasa).replace(',', '.'))
            if nomeaulaa == 'Tripla 2h':
                aula.tripla2 = aula.tripla2 + float(str(numeroaulasa).replace(',', '.'))

            simples1 = aula.simples1
            simples2 = aula.simples2
            dupla1 = aula.dupla1
            dupla2 = aula.dupla2
            tripla1 = aula.tripla1
            tripla2 = aula.tripla2

            if simples1 + simples2 <= 10:
                internosimples = valor.internosimples1a10
            else:
                if simples1 + simples2 <= 30:
                    internosimples = valor.internosimples11a30
                else:
                    if simples1 + simples2 <= 50:
                        internosimples = valor.internosimples31a50
                    else:
                        if simples1 + simples2 <= 100:
                            internosimples = valor.internosimples51a100
                        else:
                            if simples1 + simples2 <= 120:
                                internosimples = valor.internosimples101a120
                            else:
                                internosimples = valor.internosimplesacima120

            if dupla1 + dupla2 <= 60:
                internodupla = valor.internodupla1a60
            else:
                if dupla1 + dupla2 <= 119:
                    internodupla = valor.internodupla61a119
                else:
                    internodupla = valor.internoduplaacima119

            if tripla1 + tripla2 <= 60:
                internotripla = valor.internotripla1a60
            else:
                if tripla1 + tripla2 <= 119:
                    internotripla = valor.internotripla61a119
                else:
                    internotripla = valor.internotriplaacima119

            if simples1 + simples2 <= 10:
                externosimples = valor.externosimples1a10
            else:
                if simples1 + simples2 <= 30:
                    externosimples = valor.externosimples11a30
                else:
                    if simples1 + simples2 <= 50:
                        externosimples = valor.externosimples31a50
                    else:
                        if simples1 + simples2 <= 100:
                            externosimples = valor.externosimples51a100
                        else:
                            if simples1 + simples2 <= 120:
                                externosimples = valor.externosimples101a120
                            else:
                                externosimples = valor.externosimplesacima120

            if dupla1 + dupla2 <= 60:
                externodupla = valor.externodupla1a60
            else:
                if dupla1 + dupla2 <= 119:
                    externodupla = valor.externodupla61a119
                else:
                    externodupla = valor.externoduplaacima119

            if tripla1 + tripla2 <= 60:
                externotripla = valor.externotripla1a60
            else:
                if tripla1 + tripla2 <= 119:
                    externotripla = valor.externotripla61a119
                else:
                    externotripla = valor.externotriplaacima119

            internosimplesnodesc = valor.internosimplesnodesc
            internoduplanodesc = valor.internoduplanodesc
            internotriplanodesc = valor.internotriplanodesc
            externosimplesnodesc = valor.externosimplesnodesc
            externoduplanodesc = valor.externoduplanodesc
            externotriplanodesc = valor.externotriplanodesc

            if prof.tipo_personal == 'Interno':
                valor_simples = round((simples1 * internosimples) + (simples2 * 2 * internosimples), 2)
                valor_dupla = round((dupla1 * internodupla) + (dupla2 * 2 * internodupla), 2)
                valor_tripla = round((tripla1 * internotripla) + (tripla2 * 2 * internotripla), 2)
                valor_total = round(valor_simples + valor_dupla + valor_tripla, 2)
                valor_simplesnd = round((simples1 * internosimplesnodesc) + (simples2 * 2 * internosimplesnodesc), 2)
                valor_dupland = round((dupla1 * internoduplanodesc) + (dupla2 * 2 * internoduplanodesc), 2)
                valor_tripland = round((tripla1 * internotriplanodesc) + (tripla2 * 2 * internotriplanodesc), 2)
                valor_totalnd = round(valor_simplesnd + valor_dupland + valor_tripland, 2)
            else:
                valor_simples = round((simples1 * externosimples) + (simples2 * 2 * externosimples), 2)
                valor_dupla = round((dupla1 * externodupla) + (dupla2 * 2 * externodupla), 2)
                valor_tripla = round((tripla1 * externotripla) + (tripla2 * 2 * externotripla), 2)
                valor_total = round(valor_simples + valor_dupla + valor_tripla, 2)
                valor_simplesnd = round((simples1 * externosimplesnodesc) + (simples2 * 2 * externosimplesnodesc), 2)
                valor_dupland = round((dupla1 * externoduplanodesc) + (dupla2 * 2 * externoduplanodesc), 2)
                valor_tripland = round((tripla1 * externotriplanodesc) + (tripla2 * 2 * externotriplanodesc), 2)
                valor_totalnd = round(valor_simplesnd + valor_dupland + valor_tripland, 2)
            mespassado = hoje - relativedelta(months=1)
            mesp = session.query(Aulas).filter_by(mes=f'{mespassado.month}/{mespassado.year}').filter_by(personal=prof.id).first()
            if mesp:
                creditomesp = mesp.credito
                debitomesp = mesp.debito
            else:
                creditomesp = 0
                debitomesp = 0
            aula.valortotalemdia = valor_total - creditomesp + debitomesp
            aula.valortotalatraso = valor_totalnd - creditomesp + debitomesp

        if descvalor != 0:
            aula.valortotalemdia = aula.valortotalemdia - float(str(descvalor).replace(',', '.'))
            aula.valortotalatraso = aula.valortotalatraso - float(str(descvalor).replace(',', '.'))
        if acresvalor != 0:
            aula.valortotalemdia = aula.valortotalemdia + float(str(acresvalor).replace(',', '.'))
            aula.valortotalatraso = aula.valortotalatraso + float(str(acresvalor).replace(',', '.'))
        session.commit()


def suporte(assunto, mensagem):
    usuario = session.query(Usuario).filter_by(id=1).first()
    if usuario:
        email_remetente = usuario.email
        senha = usuario.senha
        # set up the SMTP server
        s = smtplib.SMTP(host=usuario.servidor, port=usuario.porta)
        s.starttls()
        s.login(email_remetente, senha)
        msg = MIMEMultipart()
        message = f'''
        Título: {assunto}\n
        Mensagem: {mensagem}\n
        '''
        # setup the parameters of the message
        msg['From'] = email_remetente
        msg['To'] = 'suportewmdapps@gmail.com'
        msg['Subject'] = "Personal Cia Athletica - Pedido de suporte"
        msg.attach(MIMEText(message, 'plain', _charset='utf-8'))
        s.send_message(msg)
        del msg
        s.quit()
    else:
        tkinter.messagebox.showinfo(title='Cadastrar e-mail remetente!', message='Antes de enviar os e-mails, cadastre '
                                                                                 'o e-mail do remetente!')


def pagamento(nome, valorcobrado, valorpago):
    personal = nome
    if valorcobrado == 1:
        pess = session.query(Personal).filter_by(nome=personal).first()
        aula = session.query(Aulas).filter_by(personal=pess.id).filter_by(mes=f'{hoje.month}/{hoje.year}').first()
        aula.valorcobrado = aula.valortotalemdia
    else:
        pess = session.query(Personal).filter_by(nome=personal).first()
        aula = session.query(Aulas).filter_by(personal=pess.id).filter_by(mes=f'{hoje.month}/{hoje.year}').first()
        aula.valorcobrado = aula.valortotalatraso
    aula.valorpago = float(str(valorpago).replace(',', '.'))
    diferenca = aula.valorcobrado - aula.valorpago

    if diferenca < 0:
        aula.credito = diferenca * (-1)
        aula.debito = 0
    else:
        aula.credito = 0
        aula.debito = diferenca
    session.commit()


def gerarecibo(nome):
    if nome:
        usuario = session.query(Usuario).filter_by(id=1).first()
        if usuario:
            carregar = Toplevel(root)
            carregar.geometry('300x100')
            carregar.title('Gerando recibo e enviando por e-mail...')
            imgs = PhotoImage(file='./static/imgs/Icone.png')
            carregar.iconphoto(False, imgs)
            ttk.Label(carregar, text='O recibo está sendo gerado e enviado...').grid(column=1, row=1, padx=20)

            pess = session.query(Personal).filter_by(nome=nome).first()
            aula = session.query(Aulas).filter_by(personal=pess.id).filter_by(mes=f'{hoje.month}/{hoje.year}').first()
            entrada = str(aula.valorpago)
            email_remetente = usuario.email
            email_destinatario = str(pess.email).split(sep=' ')[0]
            senha = usuario.senha
            # set up the SMTP server
            s = smtplib.SMTP(host=usuario.servidor, port=usuario.porta)
            s.starttls()
            s.login(email_remetente, senha)

            if ',' in entrada:
                num1, num2 = entrada.split(',')
                numero = int(num1)
                numero2 = int(num2)
                if numero2 > 0:
                    num_ptbr = num2words(numero, lang='pt-br')
                    num_ptbr2 = num2words(numero2, lang='pt-br')
                    extenso = f'{num_ptbr}'.capitalize() + f' reais e {num_ptbr2} centavos'
                else:
                    num_ptbr = num2words(numero, lang='pt-br')
                    extenso = f'{num_ptbr}'.capitalize() + ' reais'
            elif '.' in entrada:
                num1, num2 = entrada.split('.')
                numero = int(num1)
                numero2 = int(num2)
                if numero2 > 0:
                    num_ptbr = num2words(numero, lang='pt-br')
                    num_ptbr2 = num2words(numero2, lang='pt-br')
                    extenso = f'{num_ptbr}'.capitalize() + f' reais e {num_ptbr2} centavos'
                else:
                    num_ptbr = num2words(numero, lang='pt-br')
                    extenso = f'{num_ptbr}'.capitalize() + ' reais'
            else:
                num_ptbr = num2words(entrada, lang='pt-br')
                extenso = f'{num_ptbr}'.capitalize() + ' reais'
            pess.status = 'Inativo'
            aula.foipago = True
            datarec = hoje - relativedelta(months=1)
            session.commit()
            recibo = nome.split(' ')
            nomerec = str(recibo[0]).title() + ' ' + str(recibo[1]).title()
            wdFormatPDF = 17
            recibo = docx.Document('Recibo Modelo.docx')
            recibo.paragraphs[11].text = str(recibo.paragraphs[11].text).replace('#nome#', nomerec)
            recibo.paragraphs[11].text = str(recibo.paragraphs[11].text).replace('#valor#', '{:,.2f}'.format(aula.valorpago).replace('.', ','))
            recibo.paragraphs[11].text = str(recibo.paragraphs[11].text).replace('#extens', extenso)
            recibo.paragraphs[11].text = str(recibo.paragraphs[11].text).replace('#indiv#', str(aula.simples1+aula.simples2).replace('.', ','))
            recibo.paragraphs[11].text = str(recibo.paragraphs[11].text).replace('#dupl#', str(aula.dupla1+aula.dupla2).replace('.', ','))
            recibo.paragraphs[11].text = str(recibo.paragraphs[11].text).replace('#tripl#', str(aula.tripla1+aula.tripla2).replace('.', ','))
            recibo.paragraphs[11].text = str(recibo.paragraphs[11].text).replace('#mes#', str(datarec.month).zfill(2))
            recibo.paragraphs[11].text = str(recibo.paragraphs[11].text).replace('#ano#', str(datarec.year))
            recibo.paragraphs[13].text = str(recibo.paragraphs[13].text).replace('#valorcr#', '{:,.2f}'.format(aula.credito).replace('.', ','))
            recibo.paragraphs[14].text = str(recibo.paragraphs[14].text).replace('#valordeb#', '{:,.2f}'.format(aula.debito).replace('.', ','))
            recibo.save(f'.\\Relatorios\\{hoje.year}\\{str(hoje.month).zfill(2)}\\Recibos\\Recibo {nomerec}.docx')
            t.sleep(1)
            entrada = os.path.abspath(f'.\\Relatorios\\{hoje.year}\\{str(hoje.month).zfill(2)}\\Recibos\\Recibo {nomerec}.docx')
            saida = os.path.abspath(f'.\\Relatorios\\{hoje.year}\\{str(hoje.month).zfill(2)}\\Recibos\\Recibo {nomerec}.pdf')
            word = client.Dispatch('Word.Application')
            doc = word.Documents.Open(entrada)
            doc.SaveAs(saida, FileFormat=wdFormatPDF)
            doc.Close()
            word.Quit()
            os.remove(entrada)
            arquivo = saida
            msg = MIMEMultipart()
            message = f'''
            Olá, {nomerec}!\n
            \n
            Segue seu recibo de pagamento das aulas de personal na Cia Athletica.\n
            \n
            Atenciosamente,\n
            Marcelo Gonçalves
            '''
            # setup the parameters of the message
            msg['From'] = email_remetente
            msg['To'] = email_destinatario
            msg['Subject'] = "Recibo Personal"
            msg.attach(MIMEText(message, 'plain', _charset='utf-8'))

            # Anexo PNG
            arquivo_png = usuario.assinatura
            with open(arquivo_png, 'rb') as img_file:
                imagem = MIMEImage(img_file.read())
                msg.attach(imagem)

            # Anexo pdf
            part = MIMEBase('application', "octet-stream")
            part.set_payload(open(arquivo, "rb").read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', 'attachment', filename=f'Recibo personal {nomerec} {hoje.month}-{hoje.year}.pdf')
            msg.attach(part)
            s.send_message(msg)
            del msg
            s.quit()
            carregar.destroy()
        else:
            tkinter.messagebox.showinfo(title='Cadastrar e-mail remetente!',
                                        message='Antes de enviar os e-mails, cadastre '
                                                'o e-mail do remetente!')


def salvarrelatoriomes(mes):
    mesref, anoref = mes.split('/')

    wb = Workbook()
    wb.save(f'.\\Relatorios\\{anoref}\\{str(mesref).zfill(2)}\\Controle\\Relatório Personal {mesref}-{anoref}.xlsx')

    plan = l_w(os.path.abspath(f'.\\Relatorios\\{anoref}\\{str(mesref).zfill(2)}\\Controle\\Relatório Personal {mesref}-{anoref}.xlsx'), read_only=False)
    aba = plan['Sheet']
    aba['A1'].value = 'PERSONAL'
    aba.column_dimensions['A'].width = 65
    aba['B1'].value = 'VALOR COM DESCONTO'
    aba.column_dimensions['B'].width = 17
    aba['C1'].value = 'VALOR SEM DESCONTO'
    aba.column_dimensions['C'].width = 18
    aba['D1'].value = 'VALOR PAGO'
    aba.column_dimensions['D'].width = 16
    aba['G1'].value = f'Relatório Personal {mesref}/{anoref}'
    x = 2
    personal = session.query(Personal).all()
    totalemdia = 0
    totalematraso = 0
    totalpago = 0
    for pessoa in personal:
        aulasmes = session.query(Aulas).filter_by(mes=mes).filter_by(personal=pessoa.id).first()
        aba[f'A{x}'].value = pessoa.nome
        aba[f'B{x}'].value = aulasmes.valortotalemdia
        aba[f'C{x}'].value = aulasmes.valortotalatraso
        totalemdia += aulasmes.valortotalemdia
        totalematraso += aulasmes.valortotalatraso
        totalpago += aulasmes.valorpago
        if aulasmes.valorpago == 0:
            aba[f'D{x}'].value = 'NÃO PAGOU'
        else:
            aba[f'D{x}'].value = aulasmes.valorpago
        x += 1
    x += 1
    aba[f'A{x}'].value = 'TOTAL'
    aba[f'B{x}'].value = totalemdia
    aba[f'C{x}'].value = totalematraso
    aba[f'D{x}'].value = totalpago

    plan.save(f'.\\Relatorios\\{anoref}\\{str(mesref).zfill(2)}\\Controle\\Relatório Personal {mesref}-{anoref}.xlsx')


def reajuste(valor):
    num = float(str(valor).replace(',', '.'))/100
    print(num)
    valoraulas = session.query(Valores).filter_by(id=1).first()
    valoraulas.internosimplesnodesc = round(valoraulas.internosimplesnodesc + (valoraulas.internosimplesnodesc * num), 2)
    valoraulas.internoduplanodesc = round(valoraulas.internoduplanodesc + (valoraulas.internoduplanodesc * num), 2)
    valoraulas.internotriplanodesc = round(valoraulas.internotriplanodesc + (valoraulas.internotriplanodesc * num), 2)
    valoraulas.externosimplesnodesc = round(valoraulas.externosimplesnodesc + (valoraulas.externosimplesnodesc * num), 2)
    valoraulas.externoduplanodesc = round(valoraulas.externoduplanodesc + (valoraulas.externoduplanodesc * num), 2)
    valoraulas.externotriplanodesc = round(valoraulas.externotriplanodesc + (valoraulas.externotriplanodesc * num), 2)
    valoraulas.internosimples1a10 = round(valoraulas.internosimples1a10 + (valoraulas.internosimples1a10 * num), 2)
    valoraulas.internosimples11a30 = round(valoraulas.internosimples11a30 + (valoraulas.internosimples11a30 * num), 2)
    valoraulas.internosimples31a50 = round(valoraulas.internosimples31a50 + (valoraulas.internosimples31a50 * num), 2)
    valoraulas.internosimples51a100 = round(valoraulas.internosimples51a100 + (valoraulas.internosimples51a100 * num), 2)
    valoraulas.internosimples101a120 = round(valoraulas.internosimples101a120 + (valoraulas.internosimples101a120 * num), 2)
    valoraulas.internosimplesacima120 = round(valoraulas.internosimplesacima120 + (valoraulas.internosimplesacima120 * num), 2)
    valoraulas.internodupla1a60 = round(valoraulas.internodupla1a60 + (valoraulas.internodupla1a60 * num), 2)
    valoraulas.internodupla61a119 = round(valoraulas.internodupla61a119 + (valoraulas.internodupla61a119 * num), 2)
    valoraulas.internoduplaacima119 = round(valoraulas.internoduplaacima119 + (valoraulas.internoduplaacima119 * num), 2)
    valoraulas.internotripla1a60 = round(valoraulas.internotripla1a60 + (valoraulas.internotripla1a60 * num), 2)
    valoraulas.internotripla61a119 = round(valoraulas.internotripla61a119 + (valoraulas.internotripla61a119 * num), 2)
    valoraulas.internotriplaacima119 = round(valoraulas.internotriplaacima119 + (valoraulas.internotriplaacima119 * num), 2)
    valoraulas.externosimples1a10 = round(valoraulas.externosimples1a10 + (valoraulas.externosimples1a10 * num), 2)
    valoraulas.externosimples11a30 = round(valoraulas.externosimples11a30 + (valoraulas.externosimples11a30 * num), 2)
    valoraulas.externosimples31a50 = round(valoraulas.externosimples31a50 + (valoraulas.externosimples31a50 * num), 2)
    valoraulas.externosimples51a100 = round(valoraulas.externosimples51a100 + (valoraulas.externosimples51a100 * num), 2)
    valoraulas.externosimples101a120 = round(valoraulas.externosimples101a120 + (valoraulas.externosimples101a120 * num), 2)
    valoraulas.externosimplesacima120 = round(valoraulas.externosimplesacima120 + (valoraulas.externosimplesacima120 * num), 2)
    valoraulas.externodupla1a60 = round(valoraulas.externodupla1a60 + (valoraulas.externodupla1a60 * num), 2)
    valoraulas.externodupla61a119 = round(valoraulas.externodupla61a119 + (valoraulas.externodupla61a119 * num), 2)
    valoraulas.externoduplaacima119 = round(valoraulas.externoduplaacima119 + (valoraulas.externoduplaacima119 * num), 2)
    valoraulas.externotripla1a60 = round(valoraulas.externotripla1a60 + (valoraulas.externotripla1a60 * num), 2)
    valoraulas.externotripla61a119 = round(valoraulas.externotripla61a119 + (valoraulas.externotripla61a119 * num), 2)
    valoraulas.externotriplaacima119 = round(valoraulas.externotriplaacima119 + (valoraulas.externotriplaacima119 * num), 2)
    session.commit()


def cobrar():
    aulasnaopagas = session.query(Aulas).filter_by(mes=f'{hoje.month}/{hoje.year}').filter_by(foipago=False).all()
    usuario = session.query(Usuario).filter_by(id=1).first()
    if usuario:
        carregar = Toplevel(root)
        carregar.geometry('300x100')
        carregar.title('Enviando e-mails...')
        imgs = PhotoImage(file='./static/imgs/Icone.png')
        carregar.iconphoto(False, imgs)
        ttk.Label(carregar, text='Os e-mails estão sendo enviados...').grid(column=1, row=1, padx=20)
        email_remetente = usuario.email
        senha = usuario.senha
        # set up the SMTP server
        s = smtplib.SMTP(host=usuario.servidor, port=usuario.porta)
        s.starttls()
        s.login(email_remetente, senha)
        for aula in aulasnaopagas:
            personal = session.query(Personal).filter_by(id=aula.personal).first()
            nomesep = str(personal.nome).split(' ')
            nometratado = str(nomesep[0]).title() + ' ' + str(nomesep[1]).title()
            simples1 = aula.simples1
            simples2 = aula.simples2
            dupla1 = aula.dupla1
            dupla2 = aula.dupla2
            tripla1 = aula.tripla1
            tripla2 = aula.tripla2
            valor_total = round(aula.valortotalatraso, 2)

            msg = MIMEMultipart()
            message = f'''
            OBS: Caso você faça parte da Equipe E, favor desconsiderar este e-mail.\n
            \n
            Olá, {nometratado}!\n
            \n
            Não recebemos o pagamento de suas aulas de personal até o dia 10.\n
            Conforme previsto em contrato, após o dia 10 o valor das aulas perde o desconto de pontualidade.\n
            \n
            Segue valor sem desconto a ser pago por {simples1+simples2+dupla1+dupla2+tripla1+tripla2} aulas:\n
            Total de taxa: R$ {str("%.2f" % round(valor_total, 2)).replace('.',',')}.\n
            \n
            Dados para o pagamento:\n
            PIX: 03732305000186
            Banco: Itaú
            Agência: 6205
            C/C: 01588-3\n
            Assim que o pagamento for feito, favor responder esse e-mail com o comprovante bancário.
            \n
            Atenciosamente,\n
            Marcelo Gonçalves
            '''
            # setup the parameters of the message
            msg['From'] = email_remetente
            msg['To'] = str(personal.email).split(sep=' ')[0]
            msg['Subject'] = "Cobrança Personal Cia Athletica"
            msg.attach(MIMEText(message, 'plain', _charset='utf-8'))

            # Anexo PNG
            arquivo_png = usuario.assinatura
            with open(arquivo_png, 'rb') as img_file:
                imagem = MIMEImage(img_file.read())
                msg.attach(imagem)
            s.send_message(msg)
            del msg
        carregar.destroy()
        s.quit()
    else:
        tkinter.messagebox.showinfo(title='Cadastrar e-mail remetente!', message='Antes de enviar os e-mails, cadastre '
                                                                                 'o e-mail do remetente!')


def alterarpersonal(nome, tipo, wpp, email):
    if nome:
        pessoa = session.query(Personal).filter_by(nome=nome).first()
        pessoa.tipo_personal = tipo
        pessoa.whatsapp = str(wpp).replace('(', '').replace(')', '').replace('-', '').replace('.', '')
        pessoa.email = email
        session.commit()


def fecharmes(mes):
    aulasnpg = session.query(Aulas).filter_by(mes=mes).filter_by(foipago=False).all()
    if aulasnpg:
        for aula in aulasnpg:
            aula.debito = aula.valortotalatraso
            session.commit()


def janelacontato():
    j1 = Toplevel(root)
    j1.geometry('330x230')
    j1.title("Suporte")
    img = PhotoImage(file='./static/imgs/Icone.png')
    j1.iconphoto(False, img)
    ttk.Label(j1, text="Assunto").grid(column=1, row=1, sticky=W, pady=5, padx=15)
    assunto = StringVar()
    ttk.Entry(j1, width=47, textvariable=assunto).grid(column=1, row=2, sticky=W, pady=5, padx=15)
    ttk.Label(j1, text="Mensagem").grid(column=1, row=4, sticky=W, pady=5, padx=15)
    mensagem = StringVar()
    texto = Text(j1, width=35, height=5)
    texto.grid(column=1, row=5, sticky=W, pady=5, padx=15)
    
    def pegartexto():
        mens = texto.get('1.0', 'end-1c')
        return mens
    ttk.Button(j1, text="Solicitar Suporte", command=lambda: [suporte(assunto.get(), pegartexto()), tkinter.messagebox.showinfo(title='Pedido ok!', message='Pedido de suporte enviado com sucesso!'), j1.destroy()]).grid(column=1, row=6, padx=40, pady=10, sticky=E)


def porcent():
    jj2 = Toplevel(root)
    jj2.geometry('300x70')
    jj2.title("Alterar Valores de aulas")
    imgs = PhotoImage(file='./static/imgs/Icone.png')
    jj2.iconphoto(False, imgs)
    porc = StringVar()
    ttk.Label(jj2, text='Digite o percentual de aumento: ', width=38).grid(column=1, row=1, sticky=W, pady=5, padx=15)
    ttk.Entry(jj2, textvariable=porc, width=8).grid(column=1, row=1, sticky=E, pady=5, padx=15)
    ttk.Label(jj2, text='%', width=3).grid(column=2, row=1, sticky=W, pady=5, padx=15)
    ttk.Button(jj2, text='Lançar reajuste', command=lambda: [reajuste(porc.get()), tkinter.messagebox.showinfo(title='Reajuste feito!', message='Reajuste de horas lançado com sucesso!'), jj2.destroy()]).grid(column=1, row=4, sticky=E, pady=5, padx=15)


def janelavalores():
    j2 = Toplevel(root)
    j2.geometry('500x450')
    j2.title("Valores de aulas")
    imgs = PhotoImage(file='./static/imgs/Icone.png')
    j2.iconphoto(False, imgs)
    valores = session.query(Valores).filter_by(id=1).first()
    ttk.Label(j2, width=35, text=f"Valores para personais internos:").grid(column=1, row=1, sticky=W, pady=5, padx=15)
    ttk.Label(j2, text=f"Simples sem desconto: {valores.internoduplanodesc}").grid(column=1, row=2, sticky=W, pady=10, padx=15)
    ttk.Label(j2, text=f"Simples (1 a 10 aulas): {valores.internosimples1a10}").grid(column=1, row=3, sticky=W, padx=15)
    ttk.Label(j2, text=f"Simples (11 a 30 aulas): {valores.internosimples11a30}").grid(column=1, row=4, sticky=W, padx=15)
    ttk.Label(j2, text=f"Simples (31 a 50 aulas): {valores.internosimples31a50}").grid(column=1, row=5, sticky=W, padx=15)
    ttk.Label(j2, text=f"Simples (51 a 100 aulas): {valores.internosimples51a100}").grid(column=1, row=6, sticky=W, padx=15)
    ttk.Label(j2, text=f"Simples (101 a 120 aulas): {valores.internosimples101a120}").grid(column=1, row=7, sticky=W, padx=15)
    ttk.Label(j2, text=f"Simples (acima de 120 aulas): {valores.internosimplesacima120}").grid(column=1, row=8, sticky=W, padx=15)
    ttk.Label(j2, text=f"Dupla sem desconto: {valores.internoduplanodesc}").grid(column=1, row=9, sticky=W, pady=10, padx=15)
    ttk.Label(j2, text=f"Dupla (1 a 60 aulas): {valores.internodupla1a60}").grid(column=1, row=10, sticky=W, padx=15)
    ttk.Label(j2, text=f"Dupla (61 a 119 aulas): {valores.internodupla61a119}").grid(column=1, row=11, sticky=W, padx=15)
    ttk.Label(j2, text=f"Dupla (acima de 119 aulas): {valores.internoduplaacima119}").grid(column=1, row=12, sticky=W, padx=15)
    ttk.Label(j2, text=f"Tripla sem desconto: {valores.internotriplanodesc}").grid(column=1, row=13, sticky=W, pady=10, padx=15)
    ttk.Label(j2, text=f"Tripla (1 a 60 aulas): {valores.internotripla1a60}").grid(column=1, row=14, sticky=W, padx=15)
    ttk.Label(j2, text=f"Tripla (61 a 119 aulas): {valores.internotripla61a119}").grid(column=1, row=15, sticky=W, padx=15)
    ttk.Label(j2, text=f"Tripla (acima de 119 aulas): {valores.internotriplaacima119}").grid(column=1, row=16, sticky=W, padx=15)
    ttk.Label(j2, text=f"Valores para personais Externos:").grid(column=2, row=1, sticky=W, pady=5, padx=15)
    ttk.Label(j2, text=f"Simples sem desconto: {valores.externosimplesnodesc}").grid(column=2, row=2, sticky=W, pady=10, padx=15)
    ttk.Label(j2, text=f"Simples (1 a 10 aulas): {valores.externosimples1a10}").grid(column=2, row=3, sticky=W, padx=15)
    ttk.Label(j2, text=f"Simples (11 a 30 aulas): {valores.externosimples11a30}").grid(column=2, row=4, sticky=W, padx=15)
    ttk.Label(j2, text=f"Simples (31 a 50 aulas): {valores.externosimples31a50}").grid(column=2, row=5, sticky=W, padx=15)
    ttk.Label(j2, text=f"Simples (51 a 100 aulas): {valores.externosimples51a100}").grid(column=2, row=6, sticky=W, padx=15)
    ttk.Label(j2, text=f"Simples (101 a 120 aulas): {valores.externosimples101a120}").grid(column=2, row=7, sticky=W, padx=15)
    ttk.Label(j2, text=f"Simples (acima de 120 aulas): {valores.externosimplesacima120}").grid(column=2, row=8, sticky=W, padx=15)
    ttk.Label(j2, text=f"Dupla sem desconto: {valores.externoduplanodesc}").grid(column=2, row=9, sticky=W, pady=10, padx=15)
    ttk.Label(j2, text=f"Dupla (1 a 60 aulas): {valores.externodupla1a60}").grid(column=2, row=10, sticky=W, padx=15)
    ttk.Label(j2, text=f"Dupla (61 a 119 aulas): {valores.externodupla61a119}").grid(column=2, row=11, sticky=W, padx=15)
    ttk.Label(j2, text=f"Dupla (acima de 119 aulas): {valores.externoduplaacima119}").grid(column=2, row=12, sticky=W, padx=15)
    ttk.Label(j2, text=f"Tripla sem desconto: {valores.externotriplanodesc}").grid(column=2, row=13, sticky=W, pady=10, padx=15)
    ttk.Label(j2, text=f"Tripla (1 a 60 aulas): {valores.externotripla1a60}").grid(column=2, row=14, sticky=W, padx=15)
    ttk.Label(j2, text=f"Tripla (61 a 119 aulas): {valores.externotripla61a119}").grid(column=2, row=15, sticky=W, padx=15)
    ttk.Label(j2, text=f"Tripla (acima de 119 aulas): {valores.externotriplaacima119}").grid(column=2, row=16, sticky=W, padx=15)
    ttk.Button(j2, text="Alterar valores", command=porcent).grid(column=2, row=17, padx=40, pady=20, sticky=E)


def janelarelatorio():
    j3 = Toplevel(root)
    j3.geometry('490x300')
    j3.title("Relatório do mês")
    imag = PhotoImage(file='./static/imgs/Icone.png')
    j3.iconphoto(False, imag)
    refmes = StringVar()

    def mesderef(event):
        mes = event.widget.get()
        x = session.query(Aulas).filter_by(mes=mes).all()
        if x:
            pagas = len(session.query(Aulas).filter_by(mes=mes).filter_by(foipago=True).all())
            naopagas = len(session.query(Aulas).filter_by(mes=mes).filter_by(foipago=False).all())
            totalpago = 0
            x = session.query(Aulas).filter_by(mes=mes).filter_by(foipago=True).all()
            for aulapg in x:
                totalpago += aulapg.valorpago
            totaldevido = 0
            y = session.query(Aulas).filter_by(mes=mes).filter_by(foipago=False).all()
            for auladev in y:
                if hoje.day <= 10:
                    totaldevido += auladev.valortotalemdia
                else:
                    totaldevido += auladev.valortotalatraso

            tpers.config(text=f'Total de personais: {pers}')
            pagantes.config(text=f'{pagas} já efetuaram o pagamento')
            npagantes.config(text=f'{naopagas} ainda não efetuaram o pagamento')
            trec.config(text='Total recebido até o momento: R$ {:,.2f}'.format(round(totalpago,2)).replace(',','_').replace('.',',').replace('_','.'))
            tarec.config(text='Total a receber: R$ {:,.2f}'.format(round(totaldevido,2)).replace(',','_').replace('.',',').replace('_','.'))
        else:
            tpers.config(text='Total de personais: 0')
            pagantes.config(text='0 já efetuaram o pagamento')
            npagantes.config(text='0 ainda não efetuaram o pagamento')
            trec.config(text='Total recebido até o momento: R$ 0,00')
            tarec.config(text='Total a receber: R$ 0,00')

    comboref = ttk.Combobox(j3, width=8, values=referencia, textvariable=refmes, state="readonly")
    comboref.bind('<<ComboboxSelected>>', mesderef)
    comboref.grid(column=2, row=1, sticky=E, pady=1, padx=15)
    tpers = ttk.Label(j3, text='Total de personais: 0')
    tpers.grid(column=1, row=2, sticky=W, pady=20, padx=15)
    pagantes = ttk.Label(j3, text='0 já efetuaram o pagamento')
    pagantes.grid(column=1, row=3, sticky=W, padx=15)
    npagantes = ttk.Label(j3, text='0 ainda não efetuaram o pagamento')
    npagantes.grid(column=1, row=4, sticky=W, padx=15)
    trec = ttk.Label(j3, text='Total recebido até o momento: R$ 0,00')
    trec.grid(column=1, row=5, sticky=W, pady=5, padx=15)
    tarec = ttk.Label(j3, text='Total a receber: R$ 0,00')
    tarec.grid(column=1, row=6, sticky=W, padx=15)
    ttk.Button(j3, text="Fechar Mês", command=lambda: [fecharmes(refmes.get()),tkinter.messagebox.showinfo(title=f'Mês {refmes.get()} Fechado!',message='Valores não pagos lançados como débito!')]).grid(column=2, row=7, padx=40, pady=15, sticky=W)
    ttk.Button(j3, text="Salvar relatório em excel", command=lambda: [salvarrelatoriomes(refmes.get()), tkinter.messagebox.showinfo(title='Relatório Salvo!', message='Relatório salvo com sucesso!\n\nRelatório salvo em Personal > Relatórios > "ano" > "mês" > Controle.')]).grid(column=2, row=8, padx=40, pady=15, sticky=W)


def janeladadospers():
    j4 = Toplevel(root)
    j4.geometry('500x220')
    j4.title("Dados Personal")
    img = PhotoImage(file='./static/imgs/Icone.png')
    j4.iconphoto(False, img)
    j4l1 = ttk.Label(j4, text=f"Tipo: ")
    j4l1.grid(column=0, row=3, sticky=W, pady=5, padx=15)
    texttipo = StringVar()
    j4e1 = ttk.Entry(j4, width=35, textvariable=texttipo)
    j4e1.grid(column=1, row=3, sticky=W, pady=5, padx=15)
    j4l2 = ttk.Label(j4, text=f"Whatsapp: ")
    j4l2.grid(column=0, row=4, sticky=W, pady=5, padx=15)
    textwpp = StringVar()
    j4e2 = ttk.Entry(j4, width=35, textvariable=textwpp)
    j4e2.grid(column=1, row=4, sticky=W, pady=5, padx=15)
    j4l3 = ttk.Label(j4, text=f"E-mail: ")
    j4l3.grid(column=0, row=5, sticky=W, pady=5, padx=15)
    textemail = StringVar()
    j4e3 = ttk.Entry(j4, width=35, textvariable=textemail)
    j4e3.grid(column=1, row=5, sticky=W, pady=5, padx=15)
    quem = StringVar()

    def dados(event):
        nome = event.widget.get()
        if session.query(Personal).filter_by(nome=nome).first():
            prof = session.query(Personal).filter_by(nome=nome).first()
            texttipo.set(prof.tipo_personal)
            textwpp.set(prof.whatsapp)
            textemail.set(prof.email)

    ttk.Label(j4, text="Personal").grid(column=1, row=1, sticky=W, pady=5, padx=15)
    cb = ttk.Combobox(j4, width=50, values=todos, textvariable=quem)
    cb.grid(column=1, row=2, sticky=W, pady=5, padx=15)
    cb.bind('<<ComboboxSelected>>', dados)
    ttk.Button(j4, text="Salvar dados", command=lambda: [alterarpersonal(quem.get(), texttipo.get(), textwpp.get(), textemail.get()),
                                                                      tkinter.messagebox.showinfo(
                                                                          title='Dados Salvos!',
                                                                          message='Cadastro do(a) personal atualizado com sucesso!')]).grid(
        column=1, row=6, pady=15, sticky=E)


def janelaservidor():
    j5 = Toplevel(root)
    j5.geometry('440x250')
    j5.title("Configurações do servidor de e-mail")
    img = PhotoImage(file='./static/imgs/Icone.png')
    j5.iconphoto(False, img)
    ttk.Label(j5, text="").grid(column=1, row=1, sticky=W, pady=10, padx=15)
    ttk.Label(j5, text="Nome:").grid(column=1, row=2, sticky=W, pady=2, padx=15)
    nome = StringVar()
    entrynome = ttk.Entry(j5, width=40, textvariable=nome)
    entrynome.grid(column=2, row=2, sticky=W, pady=2, padx=15)
    ttk.Label(j5, text="E-mail:").grid(column=1, row=3, sticky=W, pady=2, padx=15)
    email = StringVar()
    entryemail = ttk.Entry(j5, width=40, textvariable=email)
    entryemail.grid(column=2, row=3, sticky=W, pady=2, padx=15)
    ttk.Label(j5, text="Senha:").grid(column=1, row=4, sticky=W, pady=2, padx=15)
    ttk.Label(j5, text="Servidor:").grid(column=1, row=5, sticky=W, pady=2, padx=15)

    def versenha():
        if entrysenha['show'] == "*":
            entrysenha.config(show='')
        else:
            entrysenha.config(show='*')
    senha = StringVar()
    entrysenha = ttk.Entry(j5, width=40, show="*", textvariable=senha)
    entrysenha.grid(column=2, row=4, sticky=W, pady=2, padx=15)
    ttk.Button(j5, width=4, text="Ver", command=versenha).grid(column=2, row=4, sticky=E)
    servidor = StringVar()
    entryservidor = ttk.Entry(j5, width=40, textvariable=servidor)
    entryservidor.grid(column=2, row=5, sticky=W, pady=2, padx=15)
    ttk.Label(j5, text="Porta:").grid(column=1, row=6, sticky=W, pady=2, padx=15)
    porta = StringVar()
    entryporta = ttk.Entry(j5, width=40, textvariable=porta)
    entryporta.grid(column=2, row=6, sticky=W, pady=2, padx=15)
    ttk.Label(j5, text="Imagem assinatura:").grid(column=1, row=7, sticky=W, pady=2, padx=15)
    caminhoassin = StringVar()

    def selecionarassinat():
        try:
            caminhoarq = tkinter.filedialog.askopenfilename(title='Selecione o arquivo de assinatura')
            caminhoassin.set(os.path.abspath(caminhoarq))
        except ValueError:
            pass

    btassinatura = ttk.Button(j5, text="Escolher Assinatura", command=selecionarassinat)
    btassinatura.grid(column=2, row=7, sticky=W, pady=2, padx=15)
    ttk.Button(j5, text="Salvar Dados", command=lambda: [cadastrarusuario(nome.get(),email.get(),senha.get(),
                                                                          servidor.get(),porta.get(), caminhoassin.get()),
                                                         tkinter.messagebox.showinfo(title='Cadastro efetuado!',
                                                                                     message='Dados de servidor salvos '
                                                                                             'com sucesso!'),
                                                         entrynome.delete(0, END), entryemail.delete(0, END),
                                                         entrysenha.delete(0, END),
                                                         entryservidor.delete(0, END), entryporta.delete(0, END)])\
                                                         .grid(column=2, row=8, padx=40, pady=10, sticky=E)


menubar = Menu(root)
filemenu = Menu(menubar, tearoff=0)
filemenu.add_command(label="Relatório do mês", command=janelarelatorio)
filemenu.add_command(label="Dados de personal", command=janeladadospers)
filemenu.add_command(label="Valores de Aulas", command=janelavalores)
filemenu.add_separator()
filemenu.add_command(label="Configurações de e-mail", command=janelaservidor)
filemenu.add_command(label="Solicitar Suporte", command=janelacontato)
filemenu.add_command(label="Sair", command=root.quit)
menubar.add_cascade(label="Configurações", menu=filemenu)


inicial = Frame(my_notebook, width=500, height=400)
ttk.Label(inicial, text=f"Última planilha carregada: {ultima_plan}").grid(column=1, row=1, padx=25, pady=5, sticky=W)
ttk.Label(inicial, text=f"Total de personais: {pers}").grid(column=1, row=2, padx=25, pady=5, sticky=W)
ttk.Label(inicial, text=f"Internos: {intern}").grid(column=1, row=3, padx=25, sticky=W)
ttk.Label(inicial, text=f"Externos: {extern}").grid(column=1, row=4, padx=25, sticky=W)
ttk.Label(inicial, text="Cadastrar Personal").grid(column=1, row=6, sticky=W, pady=15, padx=55)
ttk.Label(inicial, text="Nome").grid(column=1, row=7, padx=25, sticky=W)
entrada_nome = ttk.Entry(inicial, width=50)
entrada_nome.grid(column=1, row=10, padx=25, sticky=W)
ttk.Label(inicial, text="E-mail").grid(column=1, row=11, padx=25, sticky=W)
entrada_email = ttk.Entry(inicial, width=50)
entrada_email.grid(column=1, row=12, padx=25, sticky=E)
ttk.Label(inicial, text="Tel").grid(column=1, row=13, padx=25, sticky=W)
entrada_tel = ttk.Entry(inicial, width=50)
entrada_tel.grid(column=1, row=14, padx=25, sticky=E)
radio = IntVar()
c1 = Radiobutton(inicial, text='Interno', variable=radio, value=1).grid(column=1, row=15, sticky=E)
c2 = Radiobutton(inicial, text='Externo', variable=radio, value=2).grid(column=2, row=15, sticky=E)
ttk.Button(inicial, text="Cadastrar", command=lambda: [cadastrarpersonal(entrada_nome.get(), entrada_email.get(), entrada_tel.get(),radio.get()),entrada_nome.delete(0, END),entrada_email.delete(0, END),entrada_tel.delete(0, END),tkinter.messagebox.showinfo(title='Cadastro Ok!', message='Cadastro efetuado com sucesso!')]).grid(column=2, row=16, sticky=E)
inicial.pack(fill='both', expand=1)
my_notebook.add(inicial, text='Inicial')

caminho = StringVar()


def selecionar():
    try:
        caminhoarq = tkinter.filedialog.askopenfilename(title='Selecione o relatório de aulas de personal')
        caminho.set(str(caminhoarq))
    except ValueError:
        pass


calcular = Frame(my_notebook, width=500, height=400)
label1 = ttk.Label(calcular, text='Personal mês: ')
label1.grid(column=1, row=27, padx=28, pady=8, sticky=W)
label2 = ttk.Label(calcular, text='Total de personais: ')
label2.grid(column=1, row=28, padx=28, sticky=W)
label3 = ttk.Label(calcular, text='Total de aulas: ')
label3.grid(column=1, row=29, padx=28, sticky=W)
label4 = ttk.Label(calcular, text='Total de horas: ')
label4.grid(column=1, row=30, padx=28, sticky=W)
label5 = ttk.Label(calcular, text='Total aproximado de receita: ')
label5.grid(column=1, row=31, padx=28, sticky=W)


def mostrar():
    label1.config(text=f'Personal mês: {hoje.month}/{hoje.year}.')
    label2.config(text=f'Total de personais: {quantidade()}.')
    label3.config(text=f'Total de aulas: {totalaulas()}.')
    label4.config(text=f'Total de horas: {totalhoras()}.')
    label5.config(text=f'Total aproximado de receita: {totalreceita()}.')
    

ttk.Label(calcular, text=f"Total do mês: {pers} Personais").grid(column=1, row=1, padx=25, pady=5, sticky=W)
ttk.Label(calcular, text=f"{intern} Internos").grid(column=1, row=2, padx=25, sticky=W)
ttk.Label(calcular, text=f"{extern} Externos").grid(column=1, row=3, padx=25, sticky=W)

ttk.Button(calcular, text="Escolher Planilha", command=selecionar).grid(column=1, row=22, padx=20, pady=10, sticky=E)
ttk.Button(calcular, text="Calcular valores das aulas", command=lambda: [gerar_cobranca(caminho.get()), mostrar(),
                                                                         tkinter.messagebox.showinfo(title='Cálculo ok!', message='Cálculo de horas efetuado com sucesso!')]).grid(column=2, row=22, padx=20, pady=10, sticky=E)
if session.query(Aulas).filter_by(mes=f'{hoje.month}/{hoje.year}').first():
    ttk.Label(calcular, text=f'Personal mês: {hoje.month}/{hoje.year}.').grid(column=1, row=27, padx=28, pady=8, sticky=W)
    ttk.Label(calcular, text=f'Total de personais: {quantidade()}.').grid(column=1, row=28, padx=28, sticky=W)
    ttk.Label(calcular, text=f'Total de aulas: {totalaulas()}.').grid(column=1, row=29, padx=28, sticky=W)
    ttk.Label(calcular, text=f'Total de horas: {totalhoras()}.').grid(column=1, row=30, padx=28, sticky=W)
    ttk.Label(calcular, text=f'Total aproximado de receita: {totalreceita()}.').grid(column=1, row=31, padx=28, sticky=W)
calcular.pack(fill='both', expand=1)
my_notebook.add(calcular, text='Calcular')

combo = StringVar()
verificar = Frame(my_notebook, width=500, height=400)
ttk.Label(verificar, text="Nome").grid(column=1, row=1, padx=20, sticky=W)
label6 = ttk.Label(verificar, text='Aulas')
label6.grid(column=1, row=3, padx=20, sticky=W)
label7 = ttk.Label(verificar, text='0 Simples de 1h')
label7.grid(column=1, row=4, padx=20, sticky=W)
label8 = ttk.Label(verificar, text='0 Simples de 2h')
label8.grid(column=1, row=5, padx=20, sticky=W)
label9 = ttk.Label(verificar, text='0 Dupla de 1h')
label9.grid(column=1, row=6, padx=20, sticky=W)
label10 = ttk.Label(verificar, text='0 Dupla de 2h')
label10.grid(column=1, row=7, padx=20, sticky=W)
label11 = ttk.Label(verificar, text='0 Tripla de 1h')
label11.grid(column=1, row=8, padx=20, sticky=W)
label12 = ttk.Label(verificar, text='0 Tripla de 2h')
label12.grid(column=1, row=9, padx=20, sticky=W)
label13 = ttk.Label(verificar, text="Total de taxa com desconto: R$ 0,00")
label13.grid(column=1, row=10, padx=20, sticky=W)
label14 = ttk.Label(verificar, text="Total de taxa sem desconto: R$ 0,00")
label14.grid(column=1, row=11, padx=20, sticky=W)


def ver(event):
    nome = event.widget.get()
    x = session.query(Personal).filter_by(nome=nome).first()
    if x:
        aula = session.query(Aulas).filter_by(personal=x.id).filter_by(mes=f'{hoje.month}/{hoje.year}').first()
        if aula:
            simples1 = aula.simples1
            simples2 = aula.simples2
            dupla1 = aula.dupla1
            dupla2 = aula.dupla2
            tripla1 = aula.tripla1
            tripla2 = aula.tripla2
            label7.config(text=f'{simples1} Simples de 1h')
            label8.config(text=f'{simples2} Simples de 2h')
            label9.config(text=f'{dupla1} Dupla de 1h')
            label10.config(text=f'{dupla2} Dupla de 2h')
            label11.config(text=f'{tripla1} Tripla de 1h')
            label12.config(text=f'{tripla2} Tripla de 2h')
            label13.config(text='Total de taxa com desconto: R$ {:,.2f}'.format(aula.valortotalemdia).replace(',','_').replace('.',',').replace('_','.'))
            label14.config(text='Total de taxa sem desconto: R$ {:,.2f}'.format(aula.valortotalatraso).replace(',','_').replace('.',',').replace('_','.'))
        else:
            label7.config(text='0 Simples de 1h')
            label8.config(text='0 Simples de 2h')
            label9.config(text='0 Dupla de 1h')
            label10.config(text='0 Dupla de 2h')
            label11.config(text='0 Tripla de 1h')
            label12.config(text='0 Tripla de 2h')
            label13.config(text='Total de taxa com desconto: R$ 0,00')
            label14.config(text='Total de taxa sem desconto: R$ 0,00')


combo_nome = ttk.Combobox(verificar, values=values, textvariable=combo, state="readonly", width=40)
combo_nome.grid(column=1, row=2, padx=20, sticky=E)
combo_nome.bind('<<ComboboxSelected>>', ver)
ttk.Label(verificar, text="Ajustes").grid(column=2, row=1, sticky=N)
ttk.Label(verificar, text="Desconto").grid(column=2, row=2, pady=8, sticky=W)
ttk.Label(verificar, text="Aulas").grid(column=2, row=3, sticky=W)
comboaulad = StringVar()
combo_aulad = ttk.Combobox(verificar, values=nomesaulas, textvariable=comboaulad, state="readonly", width=14)
combo_aulad.grid(column=3, row=4, sticky=W)
aulas = StringVar()
descaula = ttk.Entry(verificar, width=10, textvariable=aulas)
descaula.grid(column=2, row=4, pady=5, sticky=W)
ttk.Label(verificar, text="R$").grid(column=2, row=5, sticky=W)
desconto = StringVar()
desc = ttk.Entry(verificar, width=10, textvariable=desconto)
desc.grid(column=2, row=6, sticky=W)
ttk.Label(verificar, text="Acréscimo").grid(column=2, row=7, pady=8, sticky=W)
ttk.Label(verificar, text="Aulas").grid(column=2, row=8, sticky=W)
comboaulaa = StringVar()
combo_aulaa = ttk.Combobox(verificar, values=nomesaulas, textvariable=comboaulaa, state="readonly", width=14)
combo_aulaa.grid(column=3, row=9, sticky=W)
aulasacr = StringVar()
acresaula = ttk.Entry(verificar, width=10, textvariable=aulasacr)
acresaula.grid(column=2, row=9, pady=5, sticky=W)
ttk.Label(verificar, text="R$").grid(column=2, row=10, sticky=W)
acrescimo = StringVar()
acres = ttk.Entry(verificar, width=10, textvariable=acrescimo)
acres.grid(column=2, row=11, sticky=W)
ttk.Button(verificar, text="Lançar ajuste", command=lambda: [
    ajuste(combo.get(), comboaulad.get(), int(aulas.get()), int(desconto.get()), comboaulaa.get(),
    int(aulasacr.get()), int(acrescimo.get())),
    tkinter.messagebox.showinfo(title='Lançamento ok!',message='Ajuste lançado com sucesso!'),
    desconto.set(''), acrescimo.set(''), aulas.set(''), aulasacr.set(''), descaula.delete(0, END),
    desc.delete(0, END), comboaulad.set(''), comboaulaa.set(''), acresaula.delete(0, END),
    acres.delete(0, END)
]).grid(column=3, row=13, sticky=E)
ttk.Button(verificar, text="Enviar todos e-mails", command=lambda: [mandar_email(), tkinter.messagebox.showinfo(title='Sucesso!', message='Operação efetuada com sucesso!')]).grid(column=1, row=20, padx=20, sticky=W, pady=20)
verificar.pack(fill='both', expand=1)
my_notebook.add(verificar, text='Verificar')


pagamentos = Frame(my_notebook, width=500, height=400)

ttk.Label(pagamentos, text="Nome").grid(column=1, row=1, padx=20, sticky=W)
labelpag6 = ttk.Label(pagamentos, text='Taxa')
labelpag6.grid(column=1, row=6, padx=20, sticky=W)
labelpag13 = ttk.Label(pagamentos, text="Total de taxa (com desconto): R$ 0,00")
labelpag13.grid(column=1, row=14, padx=20, sticky=W)
labelpag14 = ttk.Label(pagamentos, text="Total de taxa (sem desconto): R$ 0,00")
labelpag14.grid(column=1, row=15, padx=20, sticky=W)


def pag(event):
    nome = event.widget.get()
    x = session.query(Personal).filter_by(nome=nome).first()
    if x:
        aula = session.query(Aulas).filter_by(personal=x.id).filter_by(mes=f'{hoje.month}/{hoje.year}').first()
        if aula:
            labelpag13.config(text='Total de taxa (com desconto): R$ {:,.2f}'.format(aula.valortotalemdia).replace(',','_').replace('.',',').replace('_','.'))
            labelpag14.config(text='Total de taxa (sem desconto): R$ {:,.2f}'.format(aula.valortotalatraso).replace(',','_').replace('.',',').replace('_','.'))
        else:
            labelpag13.config(text='Total de taxa (com desconto): R$ 0,00')
            labelpag14.config(text='Total de taxa (sem desconto): R$ 0,00')


combo_nome = ttk.Combobox(pagamentos, values=values, textvariable=combo, state="readonly", width=50)
combo_nome.grid(column=1, row=2, padx=20, sticky=E)
combo_nome.bind('<<ComboboxSelected>>', pag)
labelpag = ttk.Label(pagamentos, text='Valor cobrado:')
labelpag.grid(column=1, row=16, padx=20, pady=10, sticky=W)
pgto = IntVar()
Radiobutton(pagamentos, value=1, variable=pgto, text='Com desc').grid(column=1, row=18, padx=20, sticky=W)
Radiobutton(pagamentos, value=2, variable=pgto, text='Sem desc').grid(column=1, row=19, padx=20, sticky=W)
labelpag = ttk.Label(pagamentos, text='Valor pago:')
labelpag.grid(column=1, row=20, pady=10, padx=20, sticky=W)
valorpag = StringVar()
entrada_pgto = ttk.Entry(pagamentos, textvariable=valorpag)
entrada_pgto.grid(column=1, row=21, padx=20, sticky=W)
ttk.Button(pagamentos, text="Lançar pagamento", command=lambda: [pagamento(combo_nome.get(),pgto.get(), valorpag.get()),tkinter.messagebox.showinfo(title='Lançamento ok!', message='Pagamento lançado com sucesso!'), entrada_pgto.delete(0, END)]).grid(column=2, row=22, sticky=W)
pagamentos.pack(fill='both', expand=1)
my_notebook.add(pagamentos, text='Pagamentos')

valor = StringVar()
recibos = Frame(my_notebook, width=500, height=400)
ttk.Label(recibos, text="Nome").grid(column=1, row=1, padx=20, sticky=W)
rec1 = ttk.Label(recibos, text="Aula simples: 0")
rec1.grid(column=1, row=4, padx=20, sticky=W)
rec2 = ttk.Label(recibos, text="Aula para dupla: 0")
rec2.grid(column=1, row=5, padx=20, sticky=W)
rec3 = ttk.Label(recibos, text="Aulas para trio: 0")
rec3.grid(column=1, row=6, padx=20, sticky=W)
rec4 = ttk.Label(recibos, text="Valor cobrado: R$ 0,00")
rec4.grid(column=1, row=7, padx=20, sticky=W)
rec6 = ttk.Label(recibos, text="Valor pago: R$ 0,00")
rec6.grid(column=1, row=8, padx=21, sticky=W)
rec7 = ttk.Label(recibos, text="Crédito: R$ 0,00")
rec7.grid(column=1, row=9, padx=22, sticky=W)
rec8 = ttk.Label(recibos, text="Débito: R$ 0,00")
rec8.grid(column=1, row=10, padx=23, sticky=W)


def valoresrecibo(event):
    nome = event.widget.get()
    if session.query(Personal).filter_by(nome=nome).first():
        usuario = session.query(Personal).filter_by(nome=nome).first()
        aula = session.query(Aulas).filter_by(personal=usuario.id).filter_by(mes=f'{hoje.month}/{hoje.year}').first()
        if aula:
            rec1.config(text=f'Aula simples: {aula.simples1+aula.simples2}')
            rec2.config(text=f'Aula para dupla: {aula.dupla1+aula.dupla2}')
            rec3.config(text=f'Aulas para trio: {aula.tripla1+aula.tripla2}')
            rec4.config(text='Valor cobrado: R$ {:,.2f}'.format(aula.valorcobrado).replace(',','_').replace('.',',').replace('_','.'))
            rec6.config(text='Valor pago: R$ {:,.2f}'.format(aula.valorpago).replace(',','_').replace('.',',').replace('_','.'))
            rec7.config(text='Crédito: R$ {:,.2f}'.format(aula.credito).replace(',','_').replace('.',',').replace('_','.'))
            rec8.config(text='Débito: R$ {:,.2f}'.format(aula.debito).replace(',','_').replace('.',',').replace('_','.'))
        else:
            rec1.config(text=f'Aula simples: 0')
            rec2.config(text=f'Aula para dupla: 0')
            rec3.config(text=f'Aulas para trio: 0')
            rec4.config(text='Valor cobrado: R$ 0,00')
            rec6.config(text='Valor pago: R$ 0,00')
            rec7.config(text='Crédito: R$ 0,00')
            rec8.config(text='Débito: R$ 0,00')


varcomborecibo = StringVar()
comborecibo = ttk.Combobox(recibos, values=values, textvariable=varcomborecibo, width=50)
comborecibo.grid(column=1, row=2, padx=20, sticky=E)
comborecibo.bind("<<ComboboxSelected>>", valoresrecibo)
ttk.Button(recibos, text="Enviar recibo por e-mail", command=lambda: [gerarecibo(comborecibo.get()),tkinter.messagebox.showinfo(title='Recibo ok!', message='Recibo enviado com sucesso!')]).grid(column=2, row=24, padx=20, pady=20, sticky=W)
recibos.pack(fill='both', expand=1)
my_notebook.add(recibos, text='Recibos')

inadimplentes = Frame(my_notebook, width=500, height=400)
naopagas = len(session.query(Aulas).filter_by(mes=f'{hoje.month}/{hoje.year}').filter_by(foipago=False).all())
ttk.Label(inadimplentes, text='Consulta Individual:').grid(column=1, row=1, padx=20, pady=10, sticky=W)
inad1 = ttk.Label(inadimplentes, text="Aula simples: 0")
inad1.grid(column=1, row=4, padx=20, sticky=W)
inad2 = ttk.Label(inadimplentes, text="Aula para dupla: 0")
inad2.grid(column=1, row=5, padx=20, sticky=W)
inad3 = ttk.Label(inadimplentes, text="Aulas para trio: 0")
inad3.grid(column=1, row=6, padx=20, sticky=W)
inad4 = ttk.Label(inadimplentes, text="Valor devido: R$ 0,00")
inad4.grid(column=1, row=7, padx=20, sticky=W)
inad = ttk.Label(inadimplentes, text=f'Total de inadimplentes: {naopagas}')
inad.grid(column=1, row=8, padx=20, pady=35, sticky=W)


def valoresinad(event):
    nome = event.widget.get()
    if session.query(Personal).filter_by(nome=nome).first():
        usuario = session.query(Personal).filter_by(nome=nome).first()
        aula = session.query(Aulas).filter_by(personal=usuario.id).filter_by(mes=f'{hoje.month}/{hoje.year}').first()
        if aula:
            inad1.config(text=f'Aula simples: {aula.simples1+aula.simples2}')
            inad2.config(text=f'Aula para dupla: {aula.dupla1+aula.dupla2}')
            inad3.config(text=f'Aulas para trio: {aula.tripla1+aula.tripla2}')
            inad4.config(text='Valor devido: R$ {:,.2f}'.format(aula.valortotalatraso).replace(',','_').replace('.',',').replace('_','.'))
        else:
            inad1.config(text=f'Aula simples: 0')
            inad2.config(text=f'Aula para dupla: 0')
            inad3.config(text=f'Aulas para trio: 0')
            inad4.config(text='Valor devido: R$ {:,.2f}')


varcomboinad = StringVar()
comboinad = ttk.Combobox(inadimplentes, values=values, textvariable=varcomboinad, width=50)
comboinad.grid(column=1, row=2, padx=20, sticky=E)
comboinad.bind("<<ComboboxSelected>>", valoresinad)
ttk.Button(inadimplentes, text="Enviar cobrança para todos", command=lambda: [cobrar(), tkinter.messagebox.showinfo(title='Cobrança ok!', message='Cobrança enviada com sucesso!')]).grid(column=2, row=24, padx=20, pady=20, sticky=W)

inadimplentes.pack(fill='both', expand=1)
my_notebook.add(inadimplentes, text='Inadimplentes')


root.config(menu=menubar)
root.mainloop()
